from datetime import timedelta
import whisper

source_file_name_channel = 'channels/1_channel'
print(source_file_name_channel)

# load audio and pad/trim it to fit 30 seconds
audio = whisper.load_audio(source_file_name_channel + '.wav')
audio = whisper.pad_or_trim(audio)

model = whisper.load_model("medium")
# make log-Mel spectrogram and move to the same device as the model
mel = whisper.log_mel_spectrogram(audio).to(model.device)

# detect the spoken language
_, probs = model.detect_language(mel)
print(f"Detected language: {max(probs, key=probs.get)}")
print('started...')
print()

result = model.transcribe(source_file_name_channel + '.wav',)

segments = result['segments']

text_massive = []

for segment in segments:
    startTime = str(0)+str(timedelta(seconds=int(segment['start'])))
    endTime = str(0)+str(timedelta(seconds=int(segment['end'])))
    text = segment['text']
    segmentId = segment['id']+1
    segment = f"{segmentId}. {startTime} - {endTime}\n{text[1:] if text[0] == ' ' else text}"
    #print(segment)
    text_massive.append(segment)

print()
print('Finished')

from docx import Document

# сохраняем текст с таймингом

text = text_massive

# создаем новый документ
doc = Document()

# добавляем параграф с текстом
doc.add_paragraph(source_file_name + '_' + source_file_name_channel + '_' + model_name)
for key in text:
    doc.add_paragraph(key)

# сохраняем документ
doc.save('text' + '/' + source_file_name + '_' + source_file_name_channel + '_text_timing_' + model_name + '.docx')